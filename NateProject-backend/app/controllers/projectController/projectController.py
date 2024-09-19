from fastapi import (
    APIRouter,
    UploadFile,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
)
import time
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sse_starlette.sse import EventSourceResponse
from langchain import LLMChain
from langchain_core.messages import HumanMessage, SystemMessage
import asyncio
from docx import Document
import regex as re
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
import math
from io import BytesIO
from .projectHelper import (
    get_token,
    ROOT_DIRECTORY,
    load_document,
    split_document_content,
    generate_embeddings,
    get_questionnaire_data,
    generate_openAI_embeddings,
)
from models.userModel import UserToken, UserProfile
from models.projectModel import Project, File as FileUpload, SubCategory, Category
from langchain_postgres import PGVector
from db import get_db, DATABASE_URL
from typing import List
import shutil
import json
import os
from langchain_core.prompts import ChatPromptTemplate

project_url = APIRouter()
progress = {}

cancel_events = {}


@project_url.post("/upload-file/{project_id}")
async def upload_file(
    project_id: int,
    file: List[UploadFile] = File(...),
    token: str = Depends(get_token),
    db: Session = Depends(get_db),
):
    """
    Upload a file, process it, and generate embeddings.
    Args:
        file: The file to upload.
    Returns:
        A dictionary with the filename and content type.
    """
    files = []
    upload_folder = os.path.join(ROOT_DIRECTORY, "uploads")
    os.makedirs(upload_folder, exist_ok=True)

    for doc in file:
        file_extension = doc.filename.split(".")[-1]
        uploaded_file = os.path.join(upload_folder, doc.filename)
        with open(uploaded_file, "wb+") as buffer:
            shutil.copyfileobj(doc.file, buffer)
        document = load_document(
            path=uploaded_file,
            extension=file_extension,
            metadata={"project_id": project_id},
        )
        file_db = FileUpload(
            name=doc.filename, location=uploaded_file, project_id=project_id
        )
        db.add(file_db)
        active_user = db.query(UserToken).filter(UserToken.token == token).first()
        user_profile = (
            db.query(UserProfile)
            .filter(UserProfile.user_id == active_user.user_id)
            .first()
        )
        if user_profile is None:
            db.rollback()
            raise HTTPException(status_code=404, detail="User profile not found")
        project = db.query(Project).where(Project.id == project_id).first()
        db.query()
        db.query()
        split_document = split_document_content(document)
        generate_embeddings(split_document, token, project_id, user_profile.openai_key)
        db.commit()

        file_data = {
            "name": doc.filename,
            "project_id": project_id,
            "location": uploaded_file,
            "id": file_db.id,
        }
        files.append(file_data)

    return {"files": files}


@project_url.get("/progress/{project_id}")
async def progress_endpoint(
    project_id: str,
    # request: Request
):
    """
    Provides real-time progress updates for a specific project.

    This endpoint returns a Server-Sent Events (SSE) stream that continuously
    sends progress updates for the specified project. The progress is updated
    in real-time and the stream ends when the progress reaches 100%.

    Args:
        project_id (str): The ID of the project to track progress.

    Returns:
        EventSourceResponse: A stream of JSON-encoded progress updates.
    """
    try:

        async def event_generator():
            while True:
                if project_id in progress:
                    if optimization_endpoint[project_id]["percentage"] == 100:
                        json_data = json.dumps(optimization_endpoint[project_id])
                        yield json_data
                        optimization_endpoint[project_id]["percentage"] = 0
                        optimization_endpoint[project_id]["message"] = "Finished..."
                        break
                    json_data = json.dumps(optimization_endpoint[project_id])
                    yield json_data
                await asyncio.sleep(1)

        return EventSourceResponse(event_generator(), media_type="text/event-stream")
    except Exception as e:
        print(f"Event Error : {e}")


@project_url.post("/upload-project")
def upload_project(
    name: str = Form(...),
    information: str = Form(...),
    subcategory: str = Form(...),
    file: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    token: str = Depends(get_token),
):
    """
    Upload a project and associate it with a file.
    Args:
        name: Project name.
        information: Project description.
        subcategory: Project subcategory.
        file: The file associated with the project.
    Returns:
        A dictionary with the project ID and file ID.
    """
    active_user = db.query(UserToken).filter(UserToken.token == token).first()
    if not active_user:
        raise HTTPException(status_code=401, detail="Invalid token")
    user_id = active_user.user_id
    subcategory_obj = (
        db.query(SubCategory).filter(SubCategory.name == subcategory).first()
    )
    if not subcategory_obj:
        raise HTTPException(status_code=404, detail="Subcategory not found")
    user_profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
    if user_profile is None:
        raise HTTPException(status_code=404, detail="User profile not found")
    project_db = Project(
        user_id=user_id,
        name=name,
        information=information,
        category=subcategory_obj.id,
        profile=user_profile,
    )
    db.add(project_db)
    db.commit()

    files = []
    for doc in file:
        file_extension = doc.filename.split(".")[-1]
        upload_folder = os.path.join(ROOT_DIRECTORY, "uploads")
        os.makedirs(upload_folder, exist_ok=True)
        uploaded_file = os.path.join(upload_folder, doc.filename)
        with open(uploaded_file, "wb+") as buffer:
            shutil.copyfileobj(doc.file, buffer)
        document = load_document(
            path=uploaded_file,
            extension=file_extension,
            metadata={"project_id": project_db.id},
        )
        split_document = split_document_content(document)
        generate_embeddings(
            split_document, token, project_db.id, user_profile.openai_key
        )

        file_db = FileUpload(
            name=doc.filename, location=uploaded_file, project_id=project_db.id
        )
        db.add(file_db)
        db.commit()
        db.refresh(file_db)
    db.refresh(project_db)
    files.append(file_db)
    return {"project_id": project_db.id, "file_id": file_db.id}


@project_url.get("/get-projects/")
def get_projects(token: str = Depends(get_token), db: Session = Depends(get_db)):
    """
    Retrieve all projects associated with the currently authenticated user,
    identified by the provided token.
    """
    active_user = db.query(UserToken).filter(UserToken.token == token).first()
    projects = db.query(Project).filter(Project.user_id == active_user.user_id).all()
    return projects


@project_url.get("/get-project-files/{project_id}")
def get_files(project_id, db: Session = Depends(get_db), token=Depends(get_token)):
    """
    Retrieve all files associated with a specific project ID after validating the user's token.
    Raises an HTTPException if the files or user token are not found, or if an error occurs.
    """
    try:
        files = db.query(FileUpload).filter(FileUpload.project_id == project_id).all()
        active_user = db.query(UserToken).filter(UserToken.token == token).first()

        if files is None or active_user is None:
            raise HTTPException(status_code=404, detail="Error fatching files")
        return files
    except Exception as e:
        raise HTTPException(status_code=404, detail=f"{e}")


@project_url.get("/get-categories")
def get_categories(db: Session = Depends(get_db)):
    """
    Retrieve all categories with their associated subcategories.
    """
    categories = db.query(Category).all()
    result = []
    for category in categories:
        subcategories = (
            db.query(SubCategory).filter(SubCategory.category_id == category.id).all()
        )
        result.append(
            {
                "category_id": category.id,
                "category_name": category.name,
                "subcategories": [
                    {"id": sub.id, "name": sub.name} for sub in subcategories
                ],
            }
        )
    return result


@project_url.put("/update-project-category/{project_id}")
def update_project_category(
    project_id: int,
    new_category_id: int = Form(...),
    new_subcategory_id: int = Form(None),
    db: Session = Depends(get_db),
    token: str = Depends(get_token),
):
    """
    Update the category and subcategory of a specific project.
    """
    active_user = db.query(UserToken).filter(UserToken.token == token).first()
    if not active_user:
        raise HTTPException(status_code=401, detail="Invalid token")
    project = (
        db.query(Project)
        .filter(Project.id == project_id, Project.user_id == active_user.user_id)
        .first()
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    category = db.query(Category).filter(Category.id == new_category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Category not found")

    if new_subcategory_id:
        subcategory = (
            db.query(SubCategory).filter(SubCategory.id == new_subcategory_id).first()
        )
        if not subcategory:
            raise HTTPException(status_code=404, detail="Subcategory not found")
        project.category = subcategory.id
    else:
        project.category = new_category_id

    db.commit()
    db.refresh(project)

    return {
        "project_id": project.id,
        "new_category_id": project.category,
        "message": "Project category and subcategory updated successfully",
    }


@project_url.get("/delete-project/{project_id}")
def delete_project(project_id, db: Session = Depends(get_db), token=Depends(get_token)):
    try:
        active_usre = db.query(UserToken).filter(UserToken.token == token).first()
        if not active_usre:
            raise HTTPException(status_code=401, detail="Invalid token")
        project = db.query(Project).filter(Project.id == project_id).first()
        db.delete(project)
        db.commit()
        return {"message": "Project deleted successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=404, detail=f"{e}")


optimization_endpoint = {}


def analysis_questionnaire_task(project_id, uploaded_file, user_profile, cancel_event):
    try:
        question_data = []
        in_progress = True
        yield json.dumps({"percentage": 10, "message": "Database searching..."})
        for data in get_questionnaire_data(uploaded_file, project_id, user_profile.openai_key):
            question_data.extend(data['response'])
            yield json.dumps(data['progress'])

        doc_stream = BytesIO()
        doc = Document()
        doc.add_heading("Analysis Questionnaire", 0)
        if user_profile is None:
            raise HTTPException(status_code=500, detail="No user with given token")
        if cancel_event.is_set():
            yield json.dumps({"percentage": 100, "message": "Cancelled"})
            raise HTTPException(detail="Analysis process stopped", status_code=404)
        embeddings = generate_openAI_embeddings(openai_key=user_profile.openai_key)
        VECTORSEARCH = PGVector(
            connection=DATABASE_URL,
            embeddings=embeddings,
        )
        llm = ChatOpenAI(
            model_name=user_profile.openai_model,
            api_key=user_profile.openai_key,
            temperature=0.4,
        )
        prompt_template = """Answer as short and as acurate as possible for the given content if no answer then return {no_info_response}"""
        no_info_response = "Information not found"
        template = PromptTemplate(template=prompt_template)
        yield json.dumps({"percentage": 40, "message": "Database searching..."})
        total_questions = len(question_data)
        table = doc.add_table(rows=total_questions, cols=2)
        if total_questions <= 0:
            yield json.dumps({"percentage": 100, "message": "No Question Found"})
            raise HTTPException(
                status_code=404, detail="No question found in the document"
            )
        for i, question in enumerate(question_data):
            if cancel_event.is_set():
                yield json.dumps({"percentage": 100, "message": "Cancelled"})
                break
            search_results = VECTORSEARCH.similarity_search(
                query=f"{question['question']}",
                k=3,
                filter={"project_id": f"{project_id}"},
            )
            vector_response = "".join([rec.page_content for rec in search_results])
            in_progress = False
            if search_results:
                prompt = ChatPromptTemplate.from_messages([
                    ("system", "Give shortest answer only from this Context: {context}"),
                    ("user", "{question}")
                ])
                chain = prompt | llm
                response = chain.invoke({ "context": vector_response, "question": question['question'] })
                response_text = response.content
                question_cell = table.rows[i].cells[0]
                answer_cell = table.rows[i].cells[1]
                question_paragraph = question_cell.add_paragraph()
                q_run = question_paragraph.add_run(f"{question['question']}\n")
                q_run.font.color.rgb = RGBColor(104, 109, 118) 
                answer_paragraph = answer_cell.add_paragraph()
                run = answer_paragraph.add_run(f"{response_text}\n")
                if response_text.lower() in  ["n/a", "additional information"]:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                elif re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", response_text):
                    run.font.color.rgb = RGBColor(247, 239, 229)
                elif re.search(r'\d', response_text):                    
                    run.font.color.rgb = RGBColor(124, 0, 254)
                elif re.search(r"\d{1,5}\s\w+\s\w+", response_text):
                    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                else:
                    run.font.color.rgb = RGBColor(2, 21, 38)

            progress_percentage = math.ceil(50 + (i / total_questions) * 47)
            optimization_endpoint[project_id] = {
                "percentage": progress_percentage,
                "message": "Analyzing answer...",
            }
            print(f"Percentage: {optimization_endpoint}")
            yield json.dumps(optimization_endpoint[project_id])
        upload_folder = os.path.join(ROOT_DIRECTORY, "uploads")
        os.makedirs(upload_folder, exist_ok=True)
        file = os.path.join(upload_folder, f"questionnaire_file_{project_id}.docx")
        doc.save(file)
        yield json.dumps({"percentage": 100, "message": ""})
    except Exception as e:
        optimization_endpoint[project_id] = {
            "percentage": 100,
            "message": "",
        }
        raise HTTPException(status_code=500, detail=f"Error while analyzing: {str(e)}")


@project_url.post("/analysis-questionnaire-fileupload/{project_id}")
def analysis_file_upload(project_id: str, file: UploadFile = File(...)):
    try:
        upload_folder = os.path.join(ROOT_DIRECTORY, "uploads")
        os.makedirs(upload_folder, exist_ok=True)
        uploaded_file = os.path.join(upload_folder, file.filename)

        with open(uploaded_file, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        return {"file": f"{uploaded_file}", "project_id": project_id}
    except Exception as e:

        raise HTTPException(status_code=404, detail="Error uploading file")


@project_url.get("/analysis-questionnaire/{project_id}/{file_name}")
async def analysis_questionnaire(
    project_id: str,
    file_name: str,
    db: Session = Depends(get_db),
    token: str = Query(...),
):
    """
    Uploads and analyzes a questionnaire document associated with a project.
    """
    try:
        cancel_event = asyncio.Event()
        cancel_events[project_id] = cancel_event

        def event_sender():
            yield json.dumps({"percentage": 0, "message": "Initializing..."})
            upload_folder = os.path.join(ROOT_DIRECTORY, "uploads")
            uploaded_file = os.path.join(upload_folder, file_name)
            yield json.dumps(
                {"percentage": 5, "message": "Uploading questionnaire document..."}
            )
            active_user = db.query(UserToken).filter(UserToken.token == token).first()
            user_profile = (
                db.query(UserProfile)
                .filter(UserProfile.user_id == active_user.user_id)
                .first()
            )
            yield json.dumps(
                {"percentage": 10, "message": "Uploading questionnaire document..."}
            )
            for chunk in analysis_questionnaire_task(
                project_id=project_id,
                uploaded_file=uploaded_file,
                user_profile=user_profile,
                cancel_event=cancel_event,
            ):
                if cancel_event.is_set():
                    yield json.dumps({"percentage": 100, "message": "Cancelled"})
                    break
                yield chunk
            file_path = os.path.join(
                ROOT_DIRECTORY, "uploads", f"questionnaire_file_{project_id}.docx"
            )

        return EventSourceResponse(event_sender())
    except Exception as e:
        optimization_endpoint[project_id] = {"percentage": 100, "message": f""}
        raise HTTPException(status_code=500, detail=str(e))


@project_url.get("/download-file/{project_id}")
def downloaded_file(project_id: int):
    return FileResponse(
        f"uploads/questionnaire_file_{project_id}.docx",
        filename=f"questionnaire_file_{project_id}.docx",
    )


@project_url.get("/cancel-process/{project_id}")
async def cancel_process(project_id):
    if project_id in cancel_events:
        cancel_events[project_id].set()
        return {"status ": "Progress Stopped Successfully"}
    else:
        return {"status ": "Error stopping progress"}
