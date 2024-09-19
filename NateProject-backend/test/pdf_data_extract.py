# from pypdf import PdfReader
# from langchain_openai import ChatOpenAI
# from langchain import LLMChain, PromptTemplate
# llm = ChatOpenAI(model_name="gpt-4", api_key="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")
# prompt_template = """
# You are an expert in form analysis. You will be given a section of text extracted from a PDF form. 
# Please identify if the text is a question, text field, multiple-choice question, or checkbox. 
# If it is a multiple-choice question, identify the options. 
# If it is a checkbox, identify the checkbox options.
# Here is the text:
# {text}
# Answer in the format:
# Type: [Question/Text Field/Multiple-Choice/Checkbox]
# Details: [Details based on type, e.g., options if multiple-choice or checkbox]
# """
# template = PromptTemplate(template=prompt_template, input_variables=["text"])
# chain = LLMChain(llm=llm, prompt=template)
# reader = PdfReader('test/Intake Questionnaire.pdf')
# pages = reader.pages
# text_file = open("test-result.txt", 'w+')
# print(f"Total Pages: {len(pages)}")
# for page in pages:
#     text = page.extract_text()
#     print(f"Page Text:\n{text}\n")
#     response = chain.run({"text": text})
#     print("Analysis:")
#     text_file.write(response)
#     print(response)
#     print("\n" + "-"*50 + "\n")
#_____________________________________________________________________________
from pypdf import PdfReader
from langchain_core.prompts import PromptTemplate
from langchain import LLMChain
import json
from fastapi import APIRouter, Depends, HTTPException
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.prompts import MessagesPlaceholder
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_postgres import PGVector
from docx import Document
from docx.shared import Pt
from langchain_core.messages import SystemMessage, HumanMessage
from datetime import datetime
import uuid
import os
import dotenv


# llm = ChatOpenAI(model_name="gpt-4", api_key="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", temperature=0.7)

# prompt_template = """
# You are an expert in form analysis. You will be given a section of text extracted from a PDF form.
# Keep the order of the questions, text fields.
# Please identify if the text is a question, text field, multiple-choice question, or checkbox. 
# If it is a multiple-choice question, identify the options. 
# If it is a checkbox, identify the checkbox options.
# Here is the text:
# {text}
# Return the output as an array of questions, checkboxes, multichoice questions in the following JSON format:
# {json_format}
# """
# json_format = """
# [
#     {
#         "type": "checkbox",
#         "question": "",
#         "options": ["Option 1", "Option 2", "Option 3"]
#     },
#     {
#         "type": "text field",
#         "question": "",
#         "options": []
#     },
#     {
#         "type": "multichoice",
#         "question": "",
#         "options": ["Option 1", "Option 2", "Option 3"]
#     }
#     ...
# ]
# """
# template = PromptTemplate(template=prompt_template, input_variables=["text"])
# chain = LLMChain(llm=llm, prompt=template)
# reader = PdfReader('test/Intake Questionnaire.pdf')
# pages = reader.pages
# questions = []
# print(f"Total Pages: {len(pages)}")
# for page in pages:
#     text = page.extract_text()
#     print(f"Page Text:\n{text}\n")
#     response = chain.invoke({"text": text, "json_format": json_format})
#     print("Analysis:")
#     print(response)
#     print("\n" + "-"*50 + "\n")
#     try:
#         json_output = json.loads(response['text'])
#         questions.extend(json_output)
#     except json.JSONDecodeError:
#         print("Error decoding JSON response from the model.")
# with open("test-result.json", 'w') as json_file:
#     json.dump(questions, json_file, indent=4)
#________________________________________________________________
def analyse_questionnaire():
    # Create a new Document
    doc = Document()
    doc.add_heading('Questionnaire Analysis', 0)

    # Load the question data
    with open('test-result.json', 'r') as json_file:
        question_data = json.load(json_file)
        
    user_id = 1
    api_key = "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"
    embeddings = OpenAIEmbeddings(api_key=api_key)
    VECTORSEARCH = PGVector(
        connection="postgresql://postgres:123456789@localhost:5432/NateProject",
        embeddings=embeddings,
    )
    llm = ChatOpenAI(model_name="gpt-4", api_key=api_key, temperature=0.7)
    prompt_template = """For the given content {content}, provide a response that is most suitable for
    the question "{question}". Also, remove extra words if any.
    """
    template = PromptTemplate(template=prompt_template)
    for question in question_data:
        search_results = VECTORSEARCH.similarity_search(
            query=f"{question['question']}", k=3, filter={"project_id": "41"}
        )
        vector_response = "".join([doc.page_content for doc in search_results])
        if search_results:
            chain = LLMChain(llm=llm, prompt=template)
            response = chain.run({"content": vector_response, "question": question['question']})
            p = doc.add_paragraph()
            p.add_run(f"Question: {question['question']}\n").bold = True
            p.add_run(f"Response: {response}\n")
            p.add_run("\n")
    doc.save('test-file.docx')
    print("Analysis complete. Results saved to test-file.docx.")

analyse_questionnaire()